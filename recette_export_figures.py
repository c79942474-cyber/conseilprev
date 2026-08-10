# -*- coding: utf-8 -*-
"""Emporter les calculs ET les cartes — côté modules.

CE QU'ON PROTÈGE.

1. UN DOCUMENT QUI CIRCULE SANS NOUS. Il doit porter sa source, sa licence, son
   millésime et sa PRÉCISION. La quatrième est celle qu'on oublie et c'est la
   plus utile : « classe » et « lecture graphique » ne se citent pas comme
   « exact ». Et pour l'Observatoire, une licence est plus stricte que les
   autres — CC BY-ND interdit la dérivée. Elle doit être écrite.

2. UNE FIGURE ABSENTE SE DIT. Le composeur écrit « figure non jointe » à sa
   place. La faire disparaître laisserait un dossier qui promet une carte dans
   son texte et n'en porte aucune, sans que personne sache laquelle manque.

3. ON NE PROMET PAS UNE FIGURE QUI NE PEUT PAS EXISTER. Deux clés ont été
   retirées en cours de route — l'adoption et l'empreinte — parce que ces
   panneaux ne dessinent pas de SVG. Les garder aurait produit, à chaque
   export, un manque que rien ne pouvait combler.

4. CE QUI ENTRE PAR LE GUICHET EST BORNÉ. Les figures arrivent du navigateur.
   Ce qui n'est pas du PNG est refusé, ce qui dépasse la taille est refusé, et
   une forme de données étrangère au référentiel est IGNORÉE plutôt que lue à
   moitié.
"""
import base64
import io
import os
import subprocess
import sys

DEPOT = "/home/user/conseilprev"
sys.path.insert(0, DEPOT)
os.chdir(DEPOT)

ko = 0


def ok(nom, cond, detail=""):
    global ko
    print("  %s   %s%s" % ("OK " if cond else "KO ", nom,
                           (" — " + str(detail)) if detail else ""))
    if not cond:
        ko += 1


def _refus(f):
    """Le TYPE de l'erreur levée, ou None. Un refus doit être explicite : un
    module qui rendrait un document pour un format inconnu livrerait un fichier
    que personne ne pourrait ouvrir."""
    try:
        f()
    except Exception as e:                                     # noqa: BLE001
        return type(e)
    return None


import export_dc as D                                          # noqa: E402
import export_observatoire as X                                # noqa: E402
import livrables_export as L                                   # noqa: E402
from PIL import Image                                          # noqa: E402

_b = io.BytesIO()
Image.new("RGB", (600, 360), (30, 99, 168)).save(_b, "PNG")
PNG = _b.getvalue()
PNG64 = base64.b64encode(PNG).decode()

print("\n══ 1. Le guichet des figures : ce qui entre, et ce qui est refusé ══\n")

ok("un PNG valide passe", L._figure({"figures": {"k": PNG64}}, "k") == PNG)
ok("…y compris préfixé en dataURL",
   L._figure({"figures": {"k": "data:image/png;base64," + PNG64}}, "k") == PNG)
ok("une clé absente rend None, sans lever", L._figure({"figures": {}}, "k") is None)
# CE QU'ON REFUSE. Accepter n'importe quel octet reviendrait à coller dans un
# document Word ce qu'un appelant aurait choisi d'y mettre.
ok("ce qui n'est pas du PNG est refusé",
   L._figure({"figures": {"k": base64.b64encode(b"GIF89a fake").decode()}}, "k") is None)
ok("…un JPEG aussi : le format est vérifié à l'octet, pas au nom",
   L._figure({"figures": {"k": base64.b64encode(b"\xff\xd8\xff\xe0 jpeg").decode()}},
             "k") is None)
ok("…du base64 malformé aussi",
   L._figure({"figures": {"k": "%%%pas du base64%%%"}}, "k") is None)
gros = base64.b64encode(b"\x89PNG\r\n\x1a\n" + b"0" * (L.FIGURE_MAX + 10)).decode()
ok("…et au-delà de la taille limite on REFUSE, on ne tronque pas",
   L._figure({"figures": {"k": gros}}, "k") is None,
   "%d Mio" % (L.FIGURE_MAX // (1024 * 1024)))

print("\n══ 2. Une figure absente est ÉCRITE comme absente ══\n")

md = "# T\n\n![La carte](fig:presente)\n\n![L'autre carte](fig:manquante)\n"
blocs = L._blocks(md)
figs = [b for b in blocs if b[0] == "fig"]
ok("le Markdown porte deux appels de figure", len(figs) == 2, len(figs))
ok("…chacun avec sa clé et sa légende",
   figs[0][1] == ("presente", "La carte") and figs[1][1] == ("manquante", "L'autre carte"),
   figs)
meta = {"label": "Essai", "ia": False, "figures": {"presente": PNG64}}
docx = L.build_docx(md, meta)
pdf = L.build_pdf(md, meta)
ok("le Word se construit", len(docx) > 10000, len(docx))
ok("le PDF aussi", len(pdf) > 2000, len(pdf))
# LE contrôle : la figure manquante laisse une TRACE écrite.
from docx import Document                                       # noqa: E402
doc = Document(io.BytesIO(docx))
texte = "\n".join(p.text for p in doc.paragraphs)
ok("le Word nomme la figure non jointe",
   "figure non jointe" in texte and "L'autre carte" in texte,
   [l for l in texte.split("\n") if "non jointe" in l])
imgs = [r for r in doc.part.package.parts if r.content_type == "image/png"]
ok("…et il porte bien l'image de celle qui existait", len(imgs) >= 1, len(imgs))
ok("le PDF porte une image", b"/Subtype /Image" in pdf or b"/Subtype/Image" in pdf)
# DISCRIMINATION : sans figure du tout, le document se compose quand même.
vide = L.build_docx(md, {"label": "Essai", "ia": False})
tv = "\n".join(p.text for p in Document(io.BytesIO(vide)).paragraphs)
ok("sans aucune figure, le document existe et le dit deux fois",
   tv.count("figure non jointe") == 2, tv.count("figure non jointe"))

print("\n══ 3. On ne promet pas une figure qui ne peut pas exister ══\n")

cles_obs = [c for c, _ in X.FIGURES]
cles_dc = [c for c, _ in D.FIGURES]
ok("l'Observatoire déclare trois figures", len(cles_obs) == 3, cles_obs)
ok("…et pas celle de l'adoption, panneau sans SVG",
   "adoption" not in cles_obs)
ok("le Panorama en déclare trois", len(cles_dc) == 3, cles_dc)
ok("…et pas celle de l'empreinte, panneau sans carte",
   "carte-empreinte" not in cles_dc)
# Chaque clé DÉCLARÉE doit être appelée par un Markdown, et réciproquement :
# une clé sans appel ne sert à rien, un appel sans clé ne sera jamais rempli.
md_obs = X.composer("observatoire")
appels_obs = set(b[1][0] for b in L._blocks(md_obs) if b[0] == "fig")
ok("chaque figure déclarée est appelée par le texte",
   appels_obs == set(cles_obs), (sorted(appels_obs), sorted(cles_obs)))
appels_dc = set()
for d in ("parc", "implantation"):
    appels_dc |= set(b[1][0] for b in L._blocks(D.composer(d)) if b[0] == "fig")
ok("…de même côté Panorama, pour les dossiers sans devis",
   appels_dc <= set(cles_dc) and "carte-parc" in appels_dc
   and "carte-implantation" in appels_dc, sorted(appels_dc))

print("\n══ 4. Le dossier de l'Observatoire porte ce qu'il doit porter ══\n")

s = X.sante()
ok("la santé du module est verte", not s["problemes"], s["problemes"])
ok("six chapitres", all(("## %d." % i) in md_obs for i in range(1, 7)))
for quoi, marque in X.EXIGENCES:
    ok("il porte %s" % quoi, marque in md_obs)
ok("chaque chapitre nomme sa source", md_obs.count("- **Source** —") >= 5,
   md_obs.count("- **Source** —"))
ok("…sa licence", md_obs.count("- **Licence** —") >= 5, md_obs.count("- **Licence** —"))
ok("…et sa précision", md_obs.count("- **Précision** —") >= 5,
   md_obs.count("- **Précision** —"))
ok("le crédit à reproduire est rappelé pour chacune",
   md_obs.count("- **Crédit à reproduire** —") >= 5)
ok("le document se déclare SANS rédaction par IA",
   X.MARQUE["ia"] is False and "Aucun modèle de langage" in md_obs)
# DISCRIMINATION : le garde-fou doit RÉAGIR. On retire une exigence du texte.
import re                                                       # noqa: E402
vrai = X.md_brevets
try:
    X.md_brevets = lambda d=None: "## 2. Brevets\n\nRien."
    try:
        X.composer("observatoire")
        refuse = False
    except RuntimeError as e:
        refuse = "licence sans dérivée" in str(e)
        motif = str(e)
finally:
    X.md_brevets = vrai
ok("…un dossier amputé de la licence sans dérivée est REFUSÉ, pas livré",
   refuse, locals().get("motif", "")[:90])
ok("le module est revenu à son état", not X.sante()["problemes"])

print("\n══ 5. Une forme de données étrangère est ignorée, jamais devinée ══\n")

# La page porte un modèle d'AFFICHAGE : les pays y sont une liste de points à
# projeter, pas un dictionnaire de classes. Le lire comme le référentiel
# produisait un document à moitié compris — et une erreur 503.
affichage = {"modeles": {"pays": [{"n": "France", "cl": 2}]},
             "talents": {"pays": [{"n": "Chine", "o": 26}]}}
bon, refuses = X.utilisables(affichage)
ok("le modèle d'affichage est REFUSÉ en bloc", bon == {} and len(refuses) == 2,
   (sorted(bon), sorted(refuses)))
ok("…et la composition se fait alors depuis le référentiel, sans lever",
   len(X.composer("observatoire", affichage)) > 5000)
referentiel = {"brevets": X.observatoire_ia.SEED["brevets"]}
bon2, ref2 = X.utilisables(referentiel)
ok("la forme du référentiel, elle, est acceptée",
   list(bon2) == ["brevets"] and not ref2)
ok("un corps qui n'est pas un objet ne casse rien",
   X.utilisables("des octets") == ({}, []))

print("\n══ 6. Les documents se produisent, dans les deux formats ══\n")

for fmt, tete in (("docx", b"PK"), ("pdf", b"%PDF")):
    octets, mime, nom = X.produire("observatoire", fmt, {"brevets": PNG64})
    ok("Observatoire en %s" % fmt, octets.startswith(tete) and len(octets) > 5000,
       "%d Ko · %s" % (len(octets) // 1024, nom))
    ok("…au bon type MIME", fmt in mime or "openxml" in mime, mime)
    ok("…et le nom du fichier porte le jour", nom.count("-") >= 4, nom)
for fmt in ("docx", "pdf"):
    octets, _, _ = D.produire("parc", fmt, None, {"carte-parc": PNG64})
    ok("Panorama « parc » en %s, carte jointe" % fmt, len(octets) > 5000,
       "%d Ko" % (len(octets) // 1024))
ok("un format inconnu est refusé",
   _refus(lambda: X.produire("observatoire", "odt")) is ValueError)
ok("…et un dossier inconnu aussi",
   _refus(lambda: X.composer("inexistant")) is ValueError)

print("\n══ 7. Discrimination : rien de tout cela n'existait ══\n")


def _avant(marqueur, fichier):
    hs = subprocess.check_output(
        ["git", "-C", DEPOT, "log", "-S", marqueur, "--format=%H", "--", fichier],
        text=True).split()
    ref = ("%s^" % hs[-1]) if hs else "HEAD"
    try:
        return subprocess.check_output(
            ["git", "-C", DEPOT, "show", "%s:%s" % (ref, fichier)], text=True)
    except subprocess.CalledProcessError:
        return ""            # le fichier n'existait pas du tout


av = _avant("def _figure(", "livrables_export.py")
ok("avant, l'export ne savait pas poser une image", "def _figure(" not in av)
ok("…et son analyseur Markdown ignorait les figures",
   '"fig"' not in av and "fig:" not in av)
ok("le module de l'Observatoire n'existait pas",
   _avant("def md_modeles", "export_observatoire.py") == "")
ok("…ni le sérialiseur de figures du navigateur",
   _avant("function svgEnPng", "figures_export.js") == "")
avp = _avant("carte-parc", "export_dc.py")
ok("…et le Panorama exportait ses dossiers sans aucune carte",
   "FIGURES" not in avp and "fig:" not in avp)
avapp = _avant("_figures_du_corps", "app.py")
ok("l'application n'avait pas de route pour l'Observatoire",
   "export-observatoire" not in avapp)
ok("…ni de guichet pour les figures", "_figures_du_corps" not in avapp)

print("")
print("%d contrôle(s) en échec\n" % ko if ko else "tout est vert\n")
sys.exit(1 if ko else 0)
