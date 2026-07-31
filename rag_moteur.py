# -*- coding: utf-8 -*-
"""Moteur de la base de connaissance Sentinel — recherche, extraction, diagnostic.

POURQUOI CE MODULE
La base de connaissance vivait entièrement dans app.py. Deux conséquences :
elle n'était pas testable seule, et sa recherche restait en retrait de celle du
site industriel du cabinet. Le cœur est donc extrait ici — sans Flask, sans
connexion, sans état — pour être vérifié ligne à ligne.

CE QUI CHANGE POUR LE LECTEUR D'UNE RÉPONSE
La recherche était un ESCALIER : vectorielle d'abord, et le plein-texte
seulement si la vectorielle ne rendait rien. Un document qui répondait
lexicalement mais mal sémantiquement — un numéro d'article, un nom propre, un
sigle — était donc perdu dès que la recherche vectorielle rendait cinq résultats
médiocres. Les deux moteurs sont désormais interrogés ENSEMBLE et fusionnés.
"""
import io
import os
import re
import unicodedata

VERSION = "2026-07-a"

# ═══════════════════════════════════════════════════════════════════════════
# 1. ERREURS TYPÉES
#    Un défaut d'extraction n'est pas une panne du serveur. Distinguer les deux
#    évite de rendre un 500 opaque là où le bon message est « ce PDF est une
#    image, il n'a pas de texte à extraire ».
# ═══════════════════════════════════════════════════════════════════════════

class RagErreur(Exception):
    """Erreur portant un code lisible et le statut HTTP qui lui correspond."""

    MESSAGES = {
        "extension_refusee": "Ce format de fichier n’est pas accepté.",
        "fichier_vide": "Le fichier est vide.",
        "trop_gros": "Le fichier dépasse la taille autorisée.",
        "pdf_absent": "La lecture des PDF n’est pas disponible sur ce serveur.",
        "pdf_illisible": "Ce PDF n’a pas pu être lu — il est peut-être protégé ou endommagé.",
        "pdf_sans_texte": "Ce PDF ne contient pas de texte : c’est une image scannée. "
                          "Une reconnaissance optique serait nécessaire.",
        "docx_absent": "La lecture des documents Word n’est pas disponible sur ce serveur.",
        "docx_illisible": "Ce document Word n’a pas pu être lu.",
        "xlsx_absent": "La lecture des classeurs Excel n’est pas disponible sur ce serveur.",
        "xlsx_illisible": "Ce classeur Excel n’a pas pu être lu.",
        "texte_vide": "Aucun texte n’a pu être extrait de ce fichier.",
    }

    def __init__(self, code, statut=400, detail=None):
        self.code = code
        self.statut = statut
        self.detail = detail
        super().__init__(self.MESSAGES.get(code, code))

    def message(self):
        m = self.MESSAGES.get(self.code, self.code)
        return m + (" (" + self.detail + ")" if self.detail else "")


# ═══════════════════════════════════════════════════════════════════════════
# 2. FORMATS ACCEPTÉS
#    Élargis : un cabinet reçoit des notes en Markdown, des exports CSV, des
#    journaux et des classeurs, pas seulement des PDF et des Word.
# ═══════════════════════════════════════════════════════════════════════════

FORMATS_TEXTE = ("txt", "md", "csv", "log", "json", "yaml", "yml")
FORMATS_BINAIRES = ("pdf", "docx", "xlsx", "xlsm")
EXTENSIONS = FORMATS_TEXTE + FORMATS_BINAIRES


def extension_de(nom):
    return os.path.splitext(str(nom or ""))[1].lower().lstrip(".")


def valider_extension(nom):
    ext = extension_de(nom)
    if ext not in EXTENSIONS:
        raise RagErreur("extension_refusee", 415, ext or "sans extension")
    return ext


def formats_disponibles():
    """Ce que ce serveur sait réellement lire, ici et maintenant.

    Publié tel quel : promettre un format dont la bibliothèque est absente
    produit un échec au moment de l'envoi, quand l'utilisateur a déjà attendu."""
    dispo = {e: True for e in FORMATS_TEXTE}
    for ext, mod in (("pdf", "pypdf"), ("docx", "docx"), ("xlsx", "openpyxl")):
        try:
            __import__(mod)
            dispo[ext] = True
        except Exception:                                  # noqa: BLE001
            dispo[ext] = False
    dispo["xlsm"] = dispo["xlsx"]
    return dispo


def _decoder(donnees):
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return donnees.decode(enc)
        except UnicodeDecodeError:
            continue
    return donnees.decode("utf-8", errors="replace")


def extraire_texte(nom, donnees):
    """Texte brut d'un fichier. Lève une RagErreur explicite en cas d'échec."""
    if not donnees:
        raise RagErreur("fichier_vide", 400)
    ext = valider_extension(nom)

    if ext in FORMATS_TEXTE:
        return _decoder(donnees)

    if ext == "pdf":
        try:
            from pypdf import PdfReader
        except Exception:                                  # noqa: BLE001
            raise RagErreur("pdf_absent", 503)
        try:
            lecteur = PdfReader(io.BytesIO(donnees))
            texte = "\n\n".join((p.extract_text() or "") for p in lecteur.pages)
        except Exception as e:                             # noqa: BLE001
            raise RagErreur("pdf_illisible", 422, type(e).__name__)
        # Un PDF scanné rend une chaîne vide sans lever : sans ce contrôle, le
        # document serait accepté, indexé à zéro fragment, et introuvable —
        # un échec silencieux, le pire des trois.
        if not texte.strip():
            raise RagErreur("pdf_sans_texte", 422)
        return texte

    if ext == "docx":
        try:
            import docx
        except Exception:                                  # noqa: BLE001
            raise RagErreur("docx_absent", 503)
        try:
            d = docx.Document(io.BytesIO(donnees))
            parties = [p.text for p in d.paragraphs if p.text.strip()]
            # Les tableaux Word portent souvent l'essentiel d'une note de
            # cadrage : les ignorer revient à perdre la moitié du document.
            for t in d.tables:
                for ligne in t.rows:
                    cellules = [c.text.strip() for c in ligne.cells]
                    if any(cellules):
                        parties.append(" | ".join(cellules))
            return "\n\n".join(parties)
        except Exception as e:                             # noqa: BLE001
            raise RagErreur("docx_illisible", 422, type(e).__name__)

    if ext in ("xlsx", "xlsm"):
        try:
            import openpyxl
        except Exception:                                  # noqa: BLE001
            raise RagErreur("xlsx_absent", 503)
        try:
            cl = openpyxl.load_workbook(io.BytesIO(donnees), read_only=True, data_only=True)
            parties = []
            for feuille in cl.worksheets:
                parties.append("### " + str(feuille.title))
                for ligne in feuille.iter_rows(values_only=True):
                    vals = [str(v) for v in ligne if v is not None]
                    if vals:
                        parties.append(" | ".join(vals))
            return "\n".join(parties)
        except Exception as e:                             # noqa: BLE001
            raise RagErreur("xlsx_illisible", 422, type(e).__name__)

    raise RagErreur("extension_refusee", 415, ext)


# ═══════════════════════════════════════════════════════════════════════════
# 3. DÉCOUPAGE
#    Coupe au séparateur de phrase le plus proche plutôt qu'au caractère brut :
#    un fragment tranché en plein milieu d'un mot se retrouve mal vectorisé et
#    illisible quand il est cité en réponse.
# ═══════════════════════════════════════════════════════════════════════════

TAILLE_FRAGMENT = 900
RECOUVREMENT = 150


_ESPACES = re.compile(r"[ \t]+")
_LIGNES = re.compile(r"\n{3,}")


def _normaliser(t):
    """Espaces et lignes vides ramenés à une forme unique, SUR UN FRAGMENT.

    Jamais sur le texte entier : `re.sub` sur une chaîne de 30 Mo comportant
    des millions de correspondances culmine à plus de 300 Mo de mémoire vive —
    dix fois le texte — quand les fragments produits n'en pèsent que 37. C'est
    ce détail, et lui seul, qui plafonnait la taille des documents acceptés.
    Appliqué à un fragment de 900 caractères, le même nettoyage ne coûte rien."""
    return _LIGNES.sub("\n\n", _ESPACES.sub(" ", t)).strip()


def decouper_flux(texte, taille=TAILLE_FRAGMENT, recouvrement=RECOUVREMENT):
    """Fragments rendus UN PAR UN, sans jamais construire la liste complète.

    L'appelant qui les insère au fil de l'eau ne paie donc que le fragment
    courant. `decouper` reste disponible pour les usages qui veulent la liste."""
    texte = str(texte or "")
    if not texte.strip():
        return
    if len(texte) <= taille:
        frag = _normaliser(texte)
        if frag:
            yield frag
        return

    debut, n = 0, len(texte)
    while debut < n:
        fin = min(debut + taille, n)
        if fin < n:
            # Frontière propre cherchée dans le dernier quart, SUR LE TEXTE BRUT :
            # les séparateurs qu'on y cherche (« . », saut de ligne) survivent de
            # toute façon au nettoyage, qui ne touche qu'aux espaces répétés.
            depart_fenetre = debut + int(taille * 0.75)
            fenetre = texte[depart_fenetre:fin]
            for sep in ("\n\n", ". ", ".\n", " ; ", "\n"):
                pos = fenetre.rfind(sep)
                if pos > 0:
                    fin = depart_fenetre + pos + len(sep)
                    break
        frag = _normaliser(texte[debut:fin])
        if frag:
            yield frag
        if fin >= n:
            break
        debut = max(fin - recouvrement, debut + 1)


def decouper(texte, taille=TAILLE_FRAGMENT, recouvrement=RECOUVREMENT):
    return list(decouper_flux(texte, taille, recouvrement))


# ═══════════════════════════════════════════════════════════════════════════
# 3 bis. DÉCOUPAGE D'UNE SOURCE QUI ARRIVE PAR MORCEAUX
#
#    Tout ce qui précède suppose le texte ENTIER en mémoire. Pour un document
#    de 200 Mo, cette seule hypothèse fixe la taille maximale acceptable —
#    quelles que soient les optimisations en aval.
#
#    `decouper_source` lève l'hypothèse : elle consomme une source qui arrive
#    par morceaux (blocs d'un fichier, pages d'un PDF, lignes d'un tableur) et
#    rend des fragments au fil de l'eau. À aucun instant elle ne détient plus
#    qu'un morceau et le fragment en cours. La taille du document n'entre plus
#    dans l'équation.
# ═══════════════════════════════════════════════════════════════════════════

TAILLE_BLOC = 1 << 20                       # 1 Mio lu à la fois


def _prochaine_coupe(texte, debut, taille):
    """Fin du fragment commençant à `debut` : frontière de phrase la plus
    proche dans le dernier quart, ou la taille pleine à défaut."""
    fin = min(debut + taille, len(texte))
    if fin >= len(texte):
        return fin
    depart = debut + int(taille * 0.75)
    fenetre = texte[depart:fin]
    for sep in ("\n\n", ". ", ".\n", " ; ", "\n"):
        pos = fenetre.rfind(sep)
        if pos > 0:
            return depart + pos + len(sep)
    return fin


def decouper_source(morceaux, taille=TAILLE_FRAGMENT, recouvrement=RECOUVREMENT):
    """Fragments d'une source rendue par morceaux, sans jamais tout détenir.

    `morceaux` : itérable de chaînes. Un morceau peut couper une phrase en
    deux — c'est le cas d'un bloc de fichier — et le recollage est fait ici :
    on ne coupe jamais à la frontière d'un morceau, seulement à une frontière
    de phrase, exactement comme sur un texte entier."""
    reste = ""
    for m in morceaux:
        if not m:
            continue
        reste += m
        # On ne produit que les fragments dont on est SÛR : tant qu'il ne reste
        # pas de quoi remplir un fragment entier plus sa marge, la fin pourrait
        # encore recevoir du texte et la coupe serait prématurée.
        while len(reste) >= taille * 2:
            fin = _prochaine_coupe(reste, 0, taille)
            frag = _normaliser(reste[:fin])
            if frag:
                yield frag
            reste = reste[max(fin - recouvrement, 1):]
    for frag in decouper_flux(reste, taille, recouvrement):
        yield frag


def blocs_texte(flux, taille_bloc=TAILLE_BLOC):
    """Morceaux décodés d'un fichier texte lu par blocs.

    Le décodage est INCRÉMENTAL : un caractère accentué à cheval sur deux blocs
    serait sinon décodé en deux moitiés invalides, et le document se remplirait
    de losanges noirs à l'endroit exact des coupures."""
    import codecs
    decodeur = codecs.getincrementaldecoder("utf-8")(errors="replace")
    while True:
        bloc = flux.read(taille_bloc)
        if not bloc:
            break
        m = decodeur.decode(bloc)
        if m:
            yield m
    fin = decodeur.decode(b"", True)
    if fin:
        yield fin


def morceaux_pdf(flux):
    """Pages d'un PDF, une par une. pypdf lit le fichier à la demande : le
    document n'est jamais chargé en entier."""
    try:
        from pypdf import PdfReader
    except Exception:                                          # noqa: BLE001
        raise RagErreur("pdf_absent", 503)
    try:
        lecteur = PdfReader(flux)
        pages = lecteur.pages
    except Exception as e:                                     # noqa: BLE001
        raise RagErreur("pdf_illisible", 422, type(e).__name__)
    vide = True
    for page in pages:
        try:
            t = page.extract_text() or ""
        except Exception:                                      # noqa: BLE001
            continue
        if t.strip():
            vide = False
            yield t + "\n\n"
    if vide:
        raise RagErreur("pdf_sans_texte", 422)


def morceaux_docx(flux):
    try:
        import docx
    except Exception:                                          # noqa: BLE001
        raise RagErreur("docx_absent", 503)
    try:
        d = docx.Document(flux)
    except Exception as e:                                     # noqa: BLE001
        raise RagErreur("docx_illisible", 422, type(e).__name__)
    for p in d.paragraphs:
        if p.text.strip():
            yield p.text + "\n\n"
    for t in d.tables:
        for ligne in t.rows:
            cellules = [c.text.strip() for c in ligne.cells]
            if any(cellules):
                yield " | ".join(cellules) + "\n"


def morceaux_xlsx(flux):
    try:
        import openpyxl
    except Exception:                                          # noqa: BLE001
        raise RagErreur("xlsx_absent", 503)
    try:
        cl = openpyxl.load_workbook(flux, read_only=True, data_only=True)
    except Exception as e:                                     # noqa: BLE001
        raise RagErreur("xlsx_illisible", 422, type(e).__name__)
    for feuille in cl.worksheets:
        yield "### " + str(feuille.title) + "\n"
        for ligne in feuille.iter_rows(values_only=True):
            vals = [str(v) for v in ligne if v is not None]
            if vals:
                yield " | ".join(vals) + "\n"


def morceaux_de(nom, flux):
    """Morceaux de texte d'un fichier OUVERT, selon son extension.

    Le fichier est lu depuis son flux — sur disque, tel que le serveur l'a
    reçu — et non depuis une copie en mémoire. C'est ce qui permet d'accepter
    un document dont la taille dépasse la mémoire disponible."""
    ext = valider_extension(nom)
    try:
        flux.seek(0)
    except Exception:                                          # noqa: BLE001
        pass
    if ext in FORMATS_TEXTE:
        return blocs_texte(flux)
    if ext == "pdf":
        return morceaux_pdf(flux)
    if ext == "docx":
        return morceaux_docx(flux)
    if ext in ("xlsx", "xlsm"):
        return morceaux_xlsx(flux)
    raise RagErreur("extension_refusee", 415, ext)


def fragments_de_fichier(nom, flux, taille=TAILLE_FRAGMENT, recouvrement=RECOUVREMENT):
    """Fragments prêts à indexer, depuis un fichier ouvert. Le point d'entrée
    des documents volumineux : rien n'est jamais détenu en entier."""
    return decouper_source(morceaux_de(nom, flux), taille, recouvrement)


# ═══════════════════════════════════════════════════════════════════════════
# 4. REQUÊTES
#    Le plein-texte de Postgres attend une syntaxe : lui passer des mots bruts
#    joints par « | » suffit à ramener tout document contenant N'IMPORTE LEQUEL
#    des mots, y compris « le » ou « pour ». La précision s'effondre.
# ═══════════════════════════════════════════════════════════════════════════

_MOTS_VIDES = frozenset("""
au aux avec ce ces dans de des du elle en et eux il je la le les leur lui ma
mais me meme mes moi mon ne nos notre nous on ou par pas pour qu que qui sa se
ses son sur ta te tes toi ton tu un une vos votre vous c d j l a m n s t y ete
etee etees etes etant suis es est sommes etes sont serai seras sera serons
serez seront avoir etre cette cet celui celle ceux quel quelle quels quelles
plus moins tres tout tous toute toutes autre autres comme donc alors alor
""".split())

_JETON = re.compile(r"[0-9a-zàâäéèêëîïôöùûüçñ]+", re.I)


def _sans_accents(t):
    t = unicodedata.normalize("NFD", str(t or ""))
    return "".join(c for c in t if not unicodedata.combining(c)).lower()


def termes_requete(requete, mini=2):
    """Termes utiles d'une requête, mots vides écartés, doublons supprimés.

    Les termes sont réduits à des caractères alphanumériques : ils peuvent donc
    être injectés sans risque dans une `to_tsquery`, où un caractère de
    ponctuation provoquerait une erreur de syntaxe côté base."""
    vus, out = set(), []
    for brut in _JETON.findall(str(requete or "")):
        t = _sans_accents(brut)
        if len(t) < mini or t in _MOTS_VIDES or t in vus:
            continue
        vus.add(t)
        out.append(brut.lower())
    # Une requête entièrement composée de mots vides (« pour quoi et ») ne doit
    # pas rendre une liste vide : on rend alors les jetons bruts.
    if not out:
        out = [_sans_accents(t) for t in _JETON.findall(str(requete or "")) if len(t) >= mini]
    return out


# ═══════════════════════════════════════════════════════════════════════════
# 5. FUSION DE RANGS RÉCIPROQUES
#
#    Le point central de cette reprise. Deux moteurs répondent à la même
#    question par des scores qui ne sont pas comparables : une similarité
#    cosinus vaut entre 0 et 1, un `ts_rank` vaut ce qu'il veut. Les additionner
#    n'a aucun sens ; choisir l'un et jeter l'autre en perd la moitié.
#
#    La fusion de rangs réciproques ignore les scores et ne regarde que les
#    RANGS : score(d) = Σ 1/(k + rang). Un document bien placé dans les deux
#    listes passe devant un document excellent dans une seule. Méthode éprouvée
#    en recherche d'information, et indépendante de l'échelle des moteurs.
# ═══════════════════════════════════════════════════════════════════════════

K_RRF = 60


def fusionner(listes, cle=None, k=K_RRF):
    """Fusionne des listes classées en un seul classement.

    `cle` extrait l'identité d'un résultat ; par défaut le texte du fragment.
    Les listes sont passées par ordre de priorité : à égalité, la première
    rencontrée fournit l'objet conservé."""
    if cle is None:
        def cle(r):
            return r.get("texte") if isinstance(r, dict) else r[0]
    scores, garde = {}, {}
    for liste in listes:
        for rang, r in enumerate(liste or []):
            c = cle(r)
            scores[c] = scores.get(c, 0.0) + 1.0 / (k + rang + 1)
            garde.setdefault(c, r)
    ordre = sorted(scores, key=lambda c: scores[c], reverse=True)
    sortie = []
    for c in ordre:
        r = dict(garde[c]) if isinstance(garde[c], dict) else garde[c]
        if isinstance(r, dict):
            r["score_fusion"] = round(scores[c], 5)
        sortie.append(r)
    return sortie


# ═══════════════════════════════════════════════════════════════════════════
# 6. DIAGNOSTIC
#    Une base qui ne rend rien a toujours une raison. La dire économise une
#    heure de recherche : clé d'embeddings absente, extension pgvector non
#    installée, aucun document indexé, aucun terme utile dans la requête.
# ═══════════════════════════════════════════════════════════════════════════

def diagnostiquer(etat):
    """(liste de constats, capacité de recherche effective).

    `etat` : dict attendu avec les clés `base` (un magasin est joignable),
    `base_pg` (ce magasin est PostgreSQL), `pgvector`, `embeddings`,
    `documents`, `fragments`, `fragments_vectorises`.

    `base` et `base_pg` sont distincts À DESSEIN : en développement le magasin
    est un fichier SQLite. La recherche par les mots y fonctionne — la déclarer
    « indisponible » ferait chercher une panne là où il n'y en a pas."""
    constats = []
    base = etat.get("base", etat.get("base_pg"))

    if not base:
        constats.append({"niveau": "bloquant", "cle": "base",
                         "texte": "Aucune base de données joignable : la base de "
                                  "connaissance ne peut ni stocker ni retrouver de document."})
    elif not etat.get("base_pg"):
        constats.append({"niveau": "attention", "cle": "base_locale",
                         "texte": "Magasin local SQLite, et non PostgreSQL : la recherche "
                                  "par les mots reste approchée et les vecteurs ne peuvent "
                                  "pas être stockés. C’est le mode de développement."})
    if not etat.get("documents"):
        constats.append({"niveau": "attention", "cle": "vide",
                         "texte": "Aucun document versé : la recherche ne peut rien rendre. "
                                  "C’est une base vide, pas une panne."})
    elif not etat.get("fragments"):
        constats.append({"niveau": "bloquant", "cle": "non_indexe",
                         "texte": "Des documents sont présents mais aucun n’est découpé "
                                  "en fragments : l’indexation ne s’est pas terminée."})

    if not etat.get("embeddings"):
        constats.append({"niveau": "attention", "cle": "embeddings",
                         "texte": "Clé d’API d’embeddings absente : la recherche par le SENS "
                                  "est indisponible. La recherche par les MOTS fonctionne."})
    elif not etat.get("pgvector"):
        constats.append({"niveau": "attention", "cle": "pgvector",
                         "texte": "Extension pgvector non installée sur cette base : les "
                                  "vecteurs ne peuvent pas être stockés. Recherche par les "
                                  "mots uniquement."})
    elif etat.get("fragments") and not etat.get("fragments_vectorises"):
        constats.append({"niveau": "attention", "cle": "vecteurs_absents",
                         "texte": "Les fragments existent mais aucun n’est vectorisé : "
                                  "relancer l’indexation pour activer la recherche par le sens."})

    lexical = bool(base and etat.get("fragments"))
    vectoriel = bool(lexical and etat.get("base_pg") and etat.get("pgvector")
                     and etat.get("embeddings") and etat.get("fragments_vectorises"))
    if vectoriel:
        mode = "hybride"
    elif lexical:
        mode = "lexical"
    else:
        mode = "indisponible"

    if not constats:
        constats.append({"niveau": "ok", "cle": "ok",
                         "texte": "Recherche hybride opérationnelle : les mots et le sens "
                                  "sont interrogés ensemble, puis fusionnés."})
    return constats, mode


MODES = {
    "hybride": "Mots et sens interrogés ensemble, résultats fusionnés par rang réciproque",
    "lexical": "Mots seulement — la recherche par le sens est indisponible",
    "indisponible": "Aucune recherche possible en l’état",
}


def sante(etat=None):
    d = {"version": VERSION, "formats": formats_disponibles(),
         "taille_fragment": TAILLE_FRAGMENT, "recouvrement": RECOUVREMENT}
    if etat is not None:
        constats, mode = diagnostiquer(etat)
        d["mode"] = mode
        d["mode_libelle"] = MODES.get(mode, mode)
        d["constats"] = constats
    return d
